from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Helper function to set cell borders.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def create_document():
    doc = Document()

    # --- Title Section ---
    title = doc.add_heading('Multicast NIC Hardware Filter Simulation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('CS22B1093 - Rohan G\nNetwork Systems Design - Tutorial 3')
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph('---')

    # --- Introduction ---
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        "When a computer joins a video call or streaming service, it receives data meant for a group of users rather than just itself. "
        "This is called multicast traffic. The problem is that a network might have hundreds of multicast streams flowing through it, "
        "but your computer only cares about a few of them. Processing every single packet in software would waste CPU time and slow down your system."
    )
    doc.add_paragraph(
        "To solve this, network cards (NICs) have a clever hardware trick: they use a small filter to quickly decide which packets are "
        "worth looking at. This simulation demonstrates how that filter works, including its limitations."
    )
    doc.add_paragraph('---')

    # --- Part 1 ---
    doc.add_heading('Part 1: How Multicast Addresses Work', level=1)
    
    doc.add_heading('The Problem of Addressing', level=2)
    doc.add_paragraph(
        "Regular network traffic goes from one computer to another (unicast). But multicast sends the same data to many receivers simultaneously. "
        "Think of it like a radio broadcast - the station transmits once, and everyone tuned in receives it."
    )
    p = doc.add_paragraph()
    p.add_run("Multicast uses special IP addresses in the range ").text
    p.add_run("224.0.0.0 to 239.255.255.255").bold = True
    p.add_run(". These are called Class D addresses. When a packet with such a destination arrives at a network card, "
              "the card needs to figure out if the local machine wants it.")

    doc.add_heading('Converting IP to MAC Address', level=2)
    doc.add_paragraph(
        "Ethernet networks use MAC addresses (like 01:00:5E:01:02:03) rather than IP addresses at the hardware level. "
        "So multicast IPs must be converted to multicast MACs."
    )
    doc.add_paragraph("The rule is straightforward:")
    
    # List
    doc.add_paragraph("Start with the fixed prefix 01:00:5E", style='List Bullet')
    doc.add_paragraph("Take the lower 23 bits of the IP address", style='List Bullet')
    doc.add_paragraph("Combine them", style='List Bullet')

    doc.add_paragraph("For example:")
    doc.add_paragraph("IP 239.1.2.3 → Take 1.2.3 (lower 23 bits) → MAC 01:00:5E:01:02:03", style='List Bullet')

    p = doc.add_paragraph()
    p.add_run("The catch").bold = True
    p.add_run(": We lose 5 bits of information in this conversion. This means 32 different IP addresses will map to the exact same MAC address! "
              "For instance, 224.1.2.3 and 239.129.2.3 both become 01:00:5E:01:02:03. This is an inherent limitation of the standard.")
    
    doc.add_paragraph('---')

    # --- Part 2 ---
    doc.add_heading('Part 2: CRC-32 Hashing', level=1)
    
    doc.add_heading('What is CRC?', level=2)
    doc.add_paragraph(
        "CRC stands for Cyclic Redundancy Check. Originally designed to detect errors in data transmission, it also works well as a hash function - "
        "a way to convert any input into a fixed-size number."
    )
    doc.add_paragraph(
        "The CRC-32 algorithm takes any sequence of bytes and produces a 32-bit number. The same input always gives the same output, "
        "but even a tiny change in input produces a completely different result."
    )

    doc.add_heading('How It Works (Simplified)', level=2)
    doc.add_paragraph(
        "Imagine you have a long number and you want to check if it was transmitted correctly. CRC treats your data as a giant polynomial "
        "and divides it by a special \"generator polynomial.\" The remainder of this division becomes your CRC value."
    )
    doc.add_paragraph("In practice, we use a lookup table to speed this up. Instead of doing bit-by-bit calculations:")
    doc.add_paragraph("Pre-calculate CRC values for all 256 possible byte values", style='List Bullet')
    doc.add_paragraph("For each byte of input, look up its contribution and combine with running total", style='List Bullet')
    doc.add_paragraph("Final result is the 32-bit CRC", style='List Bullet')
    
    doc.add_paragraph("The polynomial used in Ethernet is 0x04C11DB7 (or 0xEDB88320 in reversed form). This specific polynomial has been mathematically proven to catch many types of errors.")

    doc.add_heading('Why NICs Use CRC for Filtering', level=2)
    doc.add_paragraph(
        "NICs already compute CRC for every Ethernet frame to check for transmission errors. Reusing this calculation for filtering is efficient - "
        "no extra hardware needed. The CRC of the destination MAC address becomes the \"signature\" for filtering decisions."
    )
    doc.add_paragraph('---')

    # --- Part 3 ---
    doc.add_heading('Part 3: The Hardware Hash Filter', level=1)

    doc.add_heading('The Trade-off', level=2)
    doc.add_paragraph(
        "Ideally, a NIC would store a list of every multicast group the computer has joined. But memory in hardware is expensive and limited. "
        "Real NICs might need to handle dozens of groups while keeping costs down."
    )
    doc.add_paragraph(
        "The solution is a hash table - but a very small one. Typical sizes are 64 or 128 bits. Each bit position represents a \"bucket\" "
        "where one or more groups might hash to."
    )

    doc.add_heading('How It Works', level=2)
    doc.add_paragraph("1. Joining a group:", style='List Number')
    doc.add_paragraph("Convert the multicast IP to its MAC address", style='List Bullet 2')
    doc.add_paragraph("Compute CRC-32 of the MAC", style='List Bullet 2')
    doc.add_paragraph("Use the upper N bits as an index (6 bits for 64-entry table, 7 bits for 128)", style='List Bullet 2')
    doc.add_paragraph("Set that bit to 1", style='List Bullet 2')
    
    doc.add_paragraph("2. Receiving a packet:", style='List Number')
    doc.add_paragraph("Compute CRC-32 of the destination MAC", style='List Bullet 2')
    doc.add_paragraph("Extract the same upper N bits", style='List Bullet 2')
    doc.add_paragraph("If the bit is 0: drop the packet immediately (hardware decision)", style='List Bullet 2')
    doc.add_paragraph("If the bit is 1: pass to software for further checking", style='List Bullet 2')

    doc.add_heading('The False Positive Problem', level=2)
    doc.add_paragraph(
        "Since we're cramming a huge address space (millions of possible multicast MACs) into just 64 or 128 buckets, collisions are inevitable. "
        "Two completely different multicast addresses might hash to the same bucket."
    )
    doc.add_paragraph("This means:")
    doc.add_paragraph("Wanted traffic (groups we joined): Always passes the hardware filter (Check)", style='List Bullet')
    doc.add_paragraph("Unwanted traffic (groups we didn't join): Might pass if it collides with a set bit (X)", style='List Bullet')
    doc.add_paragraph(
        "When unwanted traffic sneaks through the hardware filter, it's called a false positive. The software layer catches these and drops them, "
        "but the CPU had to waste cycles processing them."
    )
    doc.add_paragraph('---')

    # --- Part 4 ---
    doc.add_heading('Part 4: Performance Metrics', level=1)
    
    doc.add_heading('Filtering Ratio', level=2)
    doc.add_paragraph("This measures how effective the hardware filter is at blocking unwanted traffic:")
    
    # Code block style
    p_code = doc.add_paragraph()
    p_code.style = 'Quote'
    run_code = p_code.add_run("Filtering Ratio = (Packets Dropped by Hardware / Total Packets) × 100%")
    run_code.font.name = 'Courier New'
    
    doc.add_paragraph("Higher is better - it means less work for the CPU.")

    doc.add_heading('False Positive Rate', level=2)
    doc.add_paragraph("This measures how often the filter lets through unwanted traffic:")

    p_code = doc.add_paragraph()
    p_code.style = 'Quote'
    run_code = p_code.add_run("False Positive Rate = (False Positives / Packets Passed to Software) × 100%")
    run_code.font.name = 'Courier New'

    doc.add_paragraph("Lower is better. A high rate means the CPU is wasting time on packets it will reject anyway.")

    doc.add_heading('Impact of Filter Size', level=2)
    
    # Table 1
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Filter Size'
    hdr_cells[1].text = 'Buckets'
    hdr_cells[2].text = 'Collision Probability'
    hdr_cells[3].text = 'Trade-off'
    
    # Bold header
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    data = [
        ['64 bits', '64', 'Higher', 'Less hardware, more false positives'],
        ['128 bits', '128', 'Medium', 'Balanced'],
        ['256 bits', '256', 'Lower', 'More hardware, fewer false positives']
    ]
    
    for row_data in data:
        row = table.add_row().cells
        for i, text in enumerate(row_data):
            row[i].text = text

    doc.add_paragraph() # Spacer
    doc.add_paragraph("With a 64-bit filter and 10 joined groups, roughly 10 buckets are set (assuming no hash collisions between our own groups). "
                      "Random unwanted traffic has about a 10/64 = 15.6% chance of hitting a set bit and becoming a false positive.")
    doc.add_paragraph('---')

    # --- Part 5 ---
    doc.add_heading('Part 5: The Complete Packet Lifecycle', level=1)
    doc.add_paragraph("When a multicast packet arrives at your network card:")

    # ASCII Diagram
    diagram = """
[Wire] → [NIC receives frame]
           ↓
       [Compute CRC of destination MAC]
           ↓
       [Check hash table bit]
           ↓
    ┌──────┴──────┐
    ↓             ↓
 [Bit = 0]    [Bit = 1]
    ↓             ↓
 [DROP]      [Pass to CPU]
 (efficient)      ↓
              [Software checks: Is this IP in our joined list?]
                  ↓
           ┌──────┴──────┐
           ↓             ↓
        [Yes]         [No]
           ↓             ↓
      [ACCEPT]      [REJECT]
      (wanted)    (false positive)
"""
    p_diag = doc.add_paragraph()
    run_diag = p_diag.add_run(diagram)
    run_diag.font.name = 'Courier New'
    run_diag.font.size = Pt(9)
    # Ensure line spacing is tight for ASCII art
    p_diag.paragraph_format.line_spacing = 1.0
    p_diag.paragraph_format.space_after = Pt(0)
    
    doc.add_paragraph('---')

    # --- Simulation Results ---
    doc.add_heading('Simulation Results', level=1)
    doc.add_paragraph("Running the simulation with 10 groups on a 64-bit filter processing 3000 packets:")
    
    doc.add_paragraph("Filtering Ratio: ~68% (hardware blocks 2/3 of traffic)", style='List Bullet')
    doc.add_paragraph("False Positive Rate: ~37% (of packets reaching software, 1/3 are unwanted)", style='List Bullet')

    doc.add_paragraph("Comparing filter sizes with identical traffic:")

    # Table 2
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Table Grid'
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'Filter'
    hdr_cells2[1].text = 'Hardware Drops'
    hdr_cells2[2].text = 'False Positives'
    hdr_cells2[3].text = 'FP Rate'

    # Bold header
    for cell in hdr_cells2:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    data2 = [
        ['64-bit', '66%', '657', '38.7%'],
        ['128-bit', '73%', '321', '23.6%'],
        ['256-bit', '76%', '170', '14.0%']
    ]

    for row_data in data2:
        row = table2.add_row().cells
        for i, text in enumerate(row_data):
            row[i].text = text

    doc.add_paragraph() # Spacer
    doc.add_paragraph("Larger filters dramatically reduce false positives at the cost of more hardware resources.")
    doc.add_paragraph('---')

    # --- Conclusion ---
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph("Hardware multicast filtering is a practical example of engineering trade-offs:")
    doc.add_paragraph("Perfect filtering would require storing all group addresses (expensive)", style='List Bullet')
    doc.add_paragraph("No filtering would burden the CPU with every multicast packet (slow)", style='List Bullet')
    doc.add_paragraph("Hash-based filtering provides a middle ground: cheap hardware, some wasted CPU cycles", style='List Bullet')
    
    doc.add_paragraph(
        "The CRC-32 algorithm, already present in NICs for error checking, enables this filtering with minimal additional circuitry. "
        "While false positives are unavoidable, the overall reduction in CPU load makes this approach worthwhile for most applications."
    )
    doc.add_paragraph('---')

    # --- Files ---
    doc.add_heading('Files', level=1)
    doc.add_paragraph("multicast_filter_simulation.py - Main simulation demonstrating all concepts", style='List Bullet')
    doc.add_paragraph("run_benchmark.py - Extended benchmarking with multiple configurations", style='List Bullet')
    doc.add_paragraph("generate_test_data.py - Test scenario generator", style='List Bullet')

    # Save
    file_path = 'Multicast_NIC_Simulation.docx'
    doc.save(file_path)
    return file_path

print(create_document())